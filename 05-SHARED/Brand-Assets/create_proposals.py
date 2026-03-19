import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h

def add_body(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1E3A5F')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')
    return table

def add_cover(doc, title, subtitle):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('THEEB ECOSYSTEM')
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    run.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.color.rgb = DARK_BLUE
    run.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GRAY
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONFIDENTIAL')
    run.font.size = Pt(11)
    run.font.color.rgb = GOLD
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('March 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    doc.add_page_break()

def add_contact(doc, tagline):
    p = doc.add_paragraph()
    run = p.add_run('Dr. Theeb Al-Theeb')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    add_body(doc, 'Founder and Group CEO, Theeb Ecosystem')
    add_body(doc, 'theeb@namamed.sa')
    add_body(doc, '+966 50 052 0242')
    add_body(doc, 'nama.med.sa')
    add_body(doc, '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tagline)
    run.font.size = Pt(13)
    run.font.color.rgb = GOLD
    run.bold = True
    run.italic = True

def add_about(doc):
    add_heading(doc, 'About the Theeb Ecosystem', 1)
    add_body(doc, 'The Theeb Ecosystem is a network of twelve specialized healthcare entities organized into four strategic groups:')
    add_table(doc,
        ['Group', 'Focus', 'Key Entities'],
        [
            ['Theeb Knowledge Group (TKG)', 'Scientific credibility and knowledge production', 'Waraqa Publishing | Nama Medical (CRO) | Nama Academy | Nama Bio | Evidence Platform | PubHelper'],
            ['Theeb Value Group (TVG)', 'Financial intelligence and value-based healthcare', 'Integra RCM | Nexus MGA | VBH Company'],
            ['Theeb Health Delivery Group (THDG)', 'Patient-facing clinical delivery', 'Welunion Clinics | Integrative Health | TAP Pharmacy'],
            ['Theeb Technology Group (TTG)', 'AI infrastructure and clinical intelligence', 'Clinical Intelligence Infrastructure (CII)'],
        ])
    add_body(doc, '')

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc

# ============================================================
# DOCUMENT 1: PHYSICIAN LEADER
# ============================================================
print('Creating Document 1: Physician Leader Partnership Proposal...')
doc1 = setup_doc()
add_cover(doc1, 'Strategic Partnership Proposal', 'Physician Career Development Program')

# Executive Summary
add_heading(doc1, '1. Executive Summary', 1)
add_body(doc1, 'The Theeb Ecosystem presents a strategic partnership opportunity for distinguished physicians who seek to transform their clinical expertise into a lasting professional legacy. This is not a service engagement. It is the construction of a comprehensive professional entity around you.')
add_body(doc1, '')
add_body(doc1, 'Saudi Arabia produces exceptional clinicians. Yet the majority face three persistent challenges that no single provider has been able to resolve: the complexity of academic promotion requirements, the absence of structured support for public knowledge contribution, and the difficulty of documenting professional impact in a way that satisfies institutional evaluation criteria.')
add_body(doc1, '')
add_body(doc1, 'The Theeb Ecosystem addresses all three simultaneously through an integrated infrastructure of twelve specialized entities spanning research, publishing, education, clinical delivery, financial intelligence, and technology. We build with you a complete professional pathway \u2014 from your first research idea to a nationally recognized career \u2014 while you continue to focus entirely on medicine.')
doc1.add_page_break()

# Challenges
add_heading(doc1, '2. Understanding Your Challenges', 1)
add_heading(doc1, '2.1 Academic Promotion', 2)
add_body(doc1, 'Publishing requirements for academic promotion continue to intensify across Saudi universities and medical institutions. Promotion committees require publications in indexed journals \u2014 both local and international \u2014 across a range of study designs. Yet the operational infrastructure required to produce this volume of scholarly output \u2014 biostatisticians, medical writers, data analysts, journal selection expertise, and submission management \u2014 is simply not available to most practicing physicians.')
add_body(doc1, '')
add_body(doc1, 'The result is a significant gap between what promotion committees demand and what physicians can realistically deliver within the constraints of clinical practice, teaching responsibilities, and administrative duties.')

add_heading(doc1, '2.2 Public Knowledge Contribution', 2)
add_body(doc1, 'Institutional evaluation frameworks increasingly assess a physician\u2019s contribution to public health education and community engagement. Social media has become the primary channel for this contribution. However, producing evidence-based educational content at the quality and frequency required for meaningful impact demands a dedicated team of medical writers, content strategists, video producers, and digital media professionals \u2014 resources that individual physicians do not have access to.')

add_heading(doc1, '2.3 Documenting Professional Impact', 2)
add_body(doc1, 'Universities, hospitals, and pharmaceutical companies evaluate physicians not only on clinical competence but on demonstrable impact: CME contributions, community engagement metrics, research output, advisory board participation, and educational product development. Most physicians perform many of these activities informally but lack a unified system to capture, document, and present this impact in the structured format that evaluation committees require.')
doc1.add_page_break()

# Six Pillars
add_heading(doc1, '3. Our Integrated Solution: Six Pillars', 1)
add_body(doc1, 'The Theeb Ecosystem does not offer isolated services. We construct a complete professional infrastructure around your expertise, organized into six interconnected pillars.')
add_body(doc1, '')

add_heading(doc1, '3.1 Pillar One: Academic Career Acceleration', 2)
add_body(doc1, 'We transform your clinical questions and research ideas into published scholarly work in the journals that matter most for your promotion \u2014 whether local or international.')
add_body(doc1, '')
add_bullet(doc1, 'Structured conversion of clinical questions into PICO-formatted research protocols with appropriate study design, sample size calculation, and methodology selection')
add_bullet(doc1, 'Dedicated biostatisticians who execute the analysis and produce publication-ready tables and figures')
add_bullet(doc1, 'Specialized medical writers who draft manuscripts in accordance with ICMJE guidelines and target journal requirements')
add_bullet(doc1, 'Strategic journal selection aligned with your specific promotion pathway \u2014 we identify the optimal journal based on specialty, impact factor, indexing status, and promotion committee preferences')
add_bullet(doc1, 'Complete submission management: cover letters, reviewer responses, revision cycles, and follow-through to acceptance')
add_bullet(doc1, 'Support across all study types: systematic reviews, meta-analyses, observational studies, case series, clinical guidelines, narrative reviews, and textbook chapters')
add_body(doc1, '')
add_body(doc1, 'Critical distinction: We do not limit you to publishing in any specific journal. We select with you the best venue for each piece of work based on your academic requirements.', bold=True)

add_heading(doc1, '3.2 Pillar Two: Evidence-Based Public Voice', 2)
add_body(doc1, 'We build and manage a digital presence that reflects your scientific authority and creates documented community impact.')
add_body(doc1, '')
add_body(doc1, 'Every piece of public content follows our Evidence-Based Content Model:')
add_table(doc1,
    ['Step', 'Component', 'Purpose'],
    [
        ['1', 'Peer-Reviewed Statistic', 'A relevant data point from current medical literature'],
        ['2', 'Medical Explanation', 'Your expert interpretation of the clinical significance'],
        ['3', 'Clinical Experience', 'A practical insight from your daily practice'],
        ['4', 'Public Health Message', 'A clear, actionable takeaway for the audience'],
    ])
add_body(doc1, '')
add_body(doc1, 'We deliver monthly impact reports documenting follower growth, engagement rates, reach statistics, and community impact metrics \u2014 all formatted for inclusion in promotion files.')

add_heading(doc1, '3.3 Pillar Three: Educational Products and Curriculum Development', 2)
add_body(doc1, 'We design, develop, and publish training curricula under your name that are accredited by SCFHS as Continuing Medical Education programs.')
add_body(doc1, '')
add_bullet(doc1, 'Collaborative curriculum design based on your expertise and identified educational gaps')
add_bullet(doc1, 'Professional writing, structuring, and formatting with learning objectives and competency frameworks')
add_bullet(doc1, 'SCFHS accreditation submission through our accredited training academy')
add_bullet(doc1, 'Delivery through our academy, your hospital, or your university \u2014 with full logistical support')
add_bullet(doc1, 'Every teaching hour automatically recorded in a digital Learning Passport synced with SCFHS')

add_heading(doc1, '3.4 Pillar Four: Virtual or Physical Clinical Practice', 2)
add_body(doc1, 'For physicians who wish to establish their own clinical practice, we build and operate the entire infrastructure. You focus exclusively on patient care.')
add_body(doc1, '')
add_bullet(doc1, 'Virtual or physical clinic setup within our national clinic network')
add_bullet(doc1, 'Complete operational management: MOH licensing, insurance credentialing, appointment systems, electronic health records, billing')
add_bullet(doc1, 'In-clinic pharmacy for direct prescription dispensing')
add_bullet(doc1, 'Dedicated administrative team handling all non-clinical operations')
add_bullet(doc1, 'Patient data (with consent) can fuel your research projects')
add_body(doc1, '')
add_body(doc1, 'Available ownership models: Fully Owned | Partnership | Fully Managed')

add_heading(doc1, '3.5 Pillar Five: Clinical Research and Trials', 2)
add_body(doc1, 'We position you as a Principal Investigator in clinical trials and real-world evidence studies.')
add_body(doc1, '')
add_bullet(doc1, 'Access to our SFDA-licensed clinical research organization (License: CT-2025-DR-0002) for Phase I-IV trial participation')
add_bullet(doc1, 'Ready-built Clinical Trial Units with trained coordinators and Ethics Committee infrastructure')
add_bullet(doc1, 'Real-World Evidence studies using Saudi population data from our clinic network')
add_bullet(doc1, 'Bioequivalence study capability through our SFDA-licensed BE center')
add_bullet(doc1, 'Direct connections to pharmaceutical company sponsors through established industry relationships')

add_heading(doc1, '3.6 Pillar Six: Professional Protection', 2)
add_bullet(doc1, 'Medical malpractice insurance bundled with 30 SCFHS-accredited CME hours annually')
add_bullet(doc1, 'Digital Learning Passport tracking all professional development and syncing with SCFHS')
add_bullet(doc1, 'Clinical evidence platform with real-time decision support: Live Guidelines, Medical Necessity Criteria, and clinical pathways')
doc1.add_page_break()

# Growth Pathway
add_heading(doc1, '4. Your Growth Pathway', 1)
add_table(doc1,
    ['Timeline', 'Milestone', 'Cumulative Impact'],
    [
        ['Month 1', 'Research plan developed; content strategy launched; Learning Passport activated', 'Foundation established'],
        ['Month 3', 'First research project in development; digital content reaching thousands; first curriculum in design', 'Momentum building'],
        ['Month 6', 'First publication accepted in indexed journal; 5,000+ followers; first CME course accredited', 'Visible results'],
        ['Year 1', '2-3 published papers; 10,000+ followers; curriculum being taught; conference keynote speaker', 'Professional entity operational'],
        ['Year 3', '8+ publications; national guideline committee member; PI in clinical trial; influential personal brand; advisory board invitations', 'Career transformation complete'],
    ])
doc1.add_page_break()

# Why Cannot Be Replicated
add_heading(doc1, '5. Why This Partnership Cannot Be Replicated', 1)
add_body(doc1, 'Any individual provider can offer one component. What no competitor can replicate is the integrated infrastructure that connects all of these into a single, compounding system.')
add_body(doc1, '')
add_body(doc1, 'The Theeb Ecosystem comprises:')
add_bullet(doc1, 'An AI-powered research factory that converts clinical questions into published evidence')
add_bullet(doc1, 'A scientific publishing house with three indexed journals')
add_bullet(doc1, 'An SCFHS-accredited training academy')
add_bullet(doc1, 'A sovereign clinical evidence platform with 250,000+ medical knowledge nodes')
add_bullet(doc1, 'A national clinic network scaling to 71 locations')
add_bullet(doc1, 'An SFDA-licensed clinical research organization')
add_bullet(doc1, 'An SFDA-licensed bioequivalence center')
add_bullet(doc1, 'A medical insurance agency')
add_bullet(doc1, 'A clinical intelligence infrastructure powered by sovereign AI')
add_body(doc1, '')
add_body(doc1, 'All of this in one ecosystem. There is nothing like it in the Middle East.', bold=True)
doc1.add_page_break()

# Engagement Model
add_heading(doc1, '6. Engagement Model', 1)
add_table(doc1,
    ['Phase', 'Activity', 'Duration'],
    [
        ['1. Initial Consultation', 'Confidential meeting to understand your academic position, promotion requirements, and professional goals', '1-2 hours'],
        ['2. Needs Assessment', 'Evaluation of which pillars are most relevant; customized plan design', '1 week'],
        ['3. Custom Partnership Plan', 'Detailed plan: research pipeline, content strategy, curriculum opportunities, clinic options, timeline', '1 week'],
        ['4. Launch', 'Partnership agreement signed; dedicated team assigned; first projects initiated', 'Month 1'],
        ['5. Quarterly Review', 'Progress reviews with milestone tracking, impact documentation, and plan adjustment', 'Ongoing'],
    ])
doc1.add_page_break()

add_about(doc1)
doc1.add_page_break()

add_heading(doc1, 'Next Step', 1)
add_body(doc1, 'Every idea in your mind deserves to become a published paper. Every clinical experience you carry deserves to reach the community. Every day without a system managing your professional growth is a missed opportunity.')
add_body(doc1, '')
add_body(doc1, 'We invite you to a confidential initial consultation to explore how this partnership can serve your specific professional goals.')
add_body(doc1, '')
add_contact(doc1, 'We build with you. We do not sell to you.')

path1 = r'C:\Users\د. ذيب\Desktop\Theeb-Ecosystem\05-SHARED\Brand-Assets\PHYSICIAN_LEADER_PARTNERSHIP_PROPOSAL.docx'
doc1.save(path1)
print(f'Document 1 saved successfully!')

# ============================================================
# DOCUMENT 2: SCIENTIFIC SOCIETY
# ============================================================
print('Creating Document 2: Scientific Society Partnership Proposal...')
doc2 = setup_doc()
add_cover(doc2, 'Strategic Infrastructure\nPartnership Proposal', 'Scientific Society Excellence Program')

add_heading(doc2, '1. Executive Summary', 1)
add_body(doc2, 'The Theeb Ecosystem invites Saudi scientific societies to a strategic infrastructure partnership that transforms their operational capabilities, accelerates their performance classification, and generates sustainable revenue \u2014 at zero upfront investment.')
add_body(doc2, '')
add_body(doc2, 'Saudi Arabia is home to over 80 registered scientific health societies. The majority operate in Category C or D \u2014 not because they lack clinical expertise, but because they lack the operational infrastructure required to score well across the eight performance evaluation categories.')
add_body(doc2, '')
add_body(doc2, 'The Theeb Ecosystem has built a complete infrastructure stack that addresses every one of these categories. We deploy this infrastructure behind the scenes, under the society\u2019s brand and authority, and share revenue generated through journal publishing, CME programs, and conference management. The society retains 40-60% of all revenue. The society retains all academic credit, authorship, and reputation.')
doc2.add_page_break()

add_heading(doc2, '2. The Evaluation Challenge', 1)
add_body(doc2, 'Scientific health societies are evaluated across eight categories with significant implications for funding, reputation, and sustainability:')
add_body(doc2, '')
add_table(doc2,
    ['Category', 'Classification', 'Reward Allocation'],
    [
        ['A (Top 25%)', '900+ points; distinguished performance', '40% of rewards budget + conference support'],
        ['B (Second 25%)', 'Good standing', '30% of rewards budget'],
        ['C (Third 25%)', 'Underperforming', '20% of rewards budget'],
        ['D (Bottom 25%)', 'At risk', '10% of rewards budget'],
        ['E (New)', 'Societies under 11 months', 'SAR 50,000 first-year grant'],
        ['F (Failed)', '2 consecutive years failed', 'Warning; potential board dissolution'],
    ])
add_body(doc2, '')
add_body(doc2, 'The total competitive scoring potential exceeds 1,950 points. The three highest-value categories \u2014 Activities and Conferences (450 points), Research and Publications (260 points), and Community Service (240 points) \u2014 total 950 points and are precisely where most societies fail. These are also precisely where the Theeb Ecosystem delivers the greatest impact.')
doc2.add_page_break()

add_heading(doc2, '3. Our Comprehensive Infrastructure Solution', 1)
add_body(doc2, 'The following sections detail our infrastructure delivery mapped to each of the eight evaluation categories.')
add_body(doc2, '')

add_heading(doc2, '3.1 Administrative and Operational Organization (150 Points)', 2)
add_bullet(doc2, 'Governance documentation support: organizational charts, committee structures, role definitions')
add_bullet(doc2, 'CPD committee operational support with detailed quarterly activity reports')
add_bullet(doc2, 'Operational planning with measurable KPIs aligned to evaluation criteria')
add_bullet(doc2, 'Semi-annual progress reports professionally prepared for submission')
add_body(doc2, 'Scoring impact: Up to 150/150 points captured.', bold=True)

add_heading(doc2, '3.2 Membership Growth and Management (400 Points)', 2)
add_bullet(doc2, 'Digital membership management platform tracking all member categories')
add_bullet(doc2, 'Member benefits: discounted CME, clinical evidence platform, and research support')
add_bullet(doc2, 'Conference attendance support for members (local and international)')
add_bullet(doc2, 'Membership growth strategy leveraging our 2,000+ physician/year training pipeline')
add_body(doc2, 'Scoring impact: Up to 400/400 points captured.', bold=True)

add_heading(doc2, '3.3 Digital Presence (90 Points)', 2)
add_bullet(doc2, 'Professionally designed bilingual website with all evaluation requirements built in')
add_bullet(doc2, 'Active social media management across Twitter plus two additional platforms')
add_bullet(doc2, 'Regular scientific content and announcements from our 250,000+ node knowledge graph')
add_bullet(doc2, 'All institutional information maintained: bylaws, strategic plan, board members, committees')
add_body(doc2, 'Scoring impact: Up to 90/90 points captured.', bold=True)

add_heading(doc2, '3.4 Activities and Conferences (450 Points)', 2)
add_body(doc2, 'This is the highest-value category and the area where most societies struggle most.')
add_bullet(doc2, 'Year-round SCFHS-accredited CME activities in your society\u2019s name \u2014 every quarter')
add_bullet(doc2, 'SCFHS Accredited Activity Provider status preparation (50 points)')
add_bullet(doc2, 'Complete annual conference management: speakers, workshops, booths, sponsorships, logistics')
add_bullet(doc2, 'CME hour generation: 30+ accredited hours annually')
add_bullet(doc2, 'International speakers through our global network (4 points each, up to 40 points)')
add_body(doc2, 'Scoring impact: Up to 450/450 points captured.', bold=True)

add_heading(doc2, '3.5 Scientific Clubs (40 Points)', 2)
add_bullet(doc2, 'Establishment and registration of scientific clubs within your society')
add_bullet(doc2, 'Dedicated digital hub for each club; activity documentation; 50+ members tracked')
add_body(doc2, 'Scoring impact: Up to 40/40 points captured.', bold=True)

add_heading(doc2, '3.6 Community Service (240 Points)', 2)
add_bullet(doc2, 'Awareness publications: infographics and educational posters (up to 40 points)')
add_bullet(doc2, 'Community magazine: annual publication registered with King Fahd National Library (15 points)')
add_bullet(doc2, 'Awareness campaigns with documented brochures and attendance (up to 50 points)')
add_bullet(doc2, 'Professional educational video production featuring your experts (20 points)')
add_bullet(doc2, 'Volunteer management on national platform (up to 100 points)')
add_body(doc2, 'Scoring impact: Up to 240/240 points captured.', bold=True)

add_heading(doc2, '3.7 Research and Scientific Publications (260 Points)', 2)
add_body(doc2, 'This is where most societies fail \u2014 and where our ecosystem delivers the greatest differentiation.', bold=True)
add_body(doc2, '')
add_bullet(doc2, 'YOUR OWN PEER-REVIEWED JOURNAL: We launch, manage, and operate a journal in your society\u2019s name with complete editorial infrastructure and Scopus/PubMed indexing pathway. Society retains 40-60% of revenue. (60 points)')
add_bullet(doc2, 'Research support via our AI-powered research factory \u2014 systematic reviews, meta-analyses, clinical studies from idea to publication (30 points)')
add_bullet(doc2, 'Research competitions for members with structured submission and judging (20 points)')
add_bullet(doc2, 'Clinical guidelines via our Live Guidelines Platform \u2014 continuously updated, accredited internally and nationally (up to 120 points)')
add_bullet(doc2, 'Practitioner publications, scientific posters, specialty books (up to 30 points)')
add_body(doc2, 'Scoring impact: Up to 260/260 points captured.', bold=True)

add_heading(doc2, '3.8 Agreements, Partnerships, and Awards (320 Points)', 2)
add_bullet(doc2, 'SCFHS framework and working agreement facilitation (20 points)')
add_bullet(doc2, 'Government entity partnerships for specialty development (30 points)')
add_bullet(doc2, 'International university and society agreements through our global network (up to 60 points)')
add_bullet(doc2, 'Local society collaborations (up to 60 points)')
add_bullet(doc2, 'Award positioning through research quality and publication output (up to 100 points)')
add_body(doc2, 'Scoring impact: Up to 320/320 points captured.', bold=True)
doc2.add_page_break()

add_heading(doc2, '4. The Path from Category C to Category A', 1)
add_table(doc2,
    ['Category', 'Max Points', 'Typical C-D Score', 'Year 1 (With Theeb)', 'Year 2 (With Theeb)'],
    [
        ['Administrative Organization', '150', '60', '120', '140'],
        ['Memberships', '400', '80', '200', '300'],
        ['Website & Social Media', '90', '25', '80', '90'],
        ['Activities & Conferences', '450', '80', '350', '420'],
        ['Scientific Clubs', '40', '10', '30', '40'],
        ['Community Service', '240', '30', '160', '220'],
        ['Research & Publications', '260', '20', '150', '230'],
        ['Agreements & Awards', '320', '40', '120', '220'],
        ['TOTAL', '1,950', '345 (Cat. D)', '1,210 (Cat. A)', '1,660 (Cat. A+)'],
    ])
doc2.add_page_break()

add_heading(doc2, '5. Revenue Model: Zero Investment, Revenue Share', 1)
add_body(doc2, 'The society invests nothing upfront. The Theeb Ecosystem invests all infrastructure. Revenue is shared transparently.')
add_body(doc2, '')
add_heading(doc2, '5.1 Revenue Share Structure', 2)
add_table(doc2,
    ['Revenue Stream', 'Society Share', 'Theeb Share'],
    [
        ['Journal Revenue (APC + advertising)', '40-60%', '40-60%'],
        ['CME Programs (fees + sponsorships)', '40-50%', '50-60%'],
        ['Conference (registration + sponsorships)', '40-50%', '50-60%'],
        ['Research Credit and Authorship', '100%', '0%'],
        ['Clinical Guideline Authorship', '100%', '0%'],
        ['Community Impact Credit', '100%', '0%'],
    ])
add_body(doc2, '')
add_heading(doc2, '5.2 Revenue Projections by Society Size', 2)
add_table(doc2,
    ['Size', 'Journal', 'CME', 'Conference', 'Total', 'Society (50%)'],
    [
        ['Small (200)', 'SAR 200K', 'SAR 300K', 'SAR 200K', 'SAR 700K', 'SAR 350K'],
        ['Medium (500)', 'SAR 500K', 'SAR 750K', 'SAR 500K', 'SAR 1.75M', 'SAR 875K'],
        ['Large (1,000+)', 'SAR 1M', 'SAR 1.5M', 'SAR 1M', 'SAR 3.5M', 'SAR 1.75M'],
    ])
doc2.add_page_break()

add_heading(doc2, '6. Implementation Timeline', 1)
add_table(doc2,
    ['Phase', 'Timeline', 'Deliverables'],
    [
        ['Agreement', 'Month 0', 'Partnership signed; zero cost to society'],
        ['Deployment', 'Month 1-3', 'Journal launched; website live; social media active; first CME delivered'],
        ['Pipeline', 'Month 3-6', 'Research projects underway; conference planning; sponsorship sales'],
        ['Full Year', 'Month 6-12', 'Journal published; 4+ CME activities; conference held; research published'],
        ['Category A', 'Month 12', 'Evaluation with projected 900+ points'],
        ['Growth', 'Year 2', 'Journal indexing advancing; revenue flowing; international partnerships'],
        ['Leadership', 'Year 3', 'National reference; indexed journal; guidelines cited; must-attend conference'],
    ])
doc2.add_page_break()

add_heading(doc2, '7. Why This Partnership Cannot Be Replicated', 1)
add_body(doc2, 'Delivering across all eight evaluation categories simultaneously requires infrastructure spanning research, publishing, education, digital media, event management, community engagement, and institutional partnerships. No single provider in the region possesses this breadth.')
add_body(doc2, '')
add_bullet(doc2, 'A scientific publishing house with three indexed journals')
add_bullet(doc2, 'An AI-powered research factory producing systematic reviews and clinical guidelines')
add_bullet(doc2, 'An SCFHS-accredited training academy')
add_bullet(doc2, 'A Live Guidelines Platform continuously updating clinical guidelines')
add_bullet(doc2, 'A clinical evidence platform with 250,000+ medical knowledge nodes')
add_bullet(doc2, 'A national clinic network providing real-world patient data')
add_bullet(doc2, 'An SFDA-licensed clinical research organization')
add_bullet(doc2, 'A professional content production team')
add_body(doc2, '')
add_body(doc2, 'Building this ecosystem required years. Replicating it would require simultaneous construction across research, publishing, education, clinical delivery, and technology.', bold=True)
doc2.add_page_break()

add_heading(doc2, '8. Engagement Model', 1)
add_table(doc2,
    ['Phase', 'Activity', 'Duration'],
    [
        ['1. Inquiry', 'Initial conversation with society leadership', '1-2 meetings'],
        ['2. Assessment', 'Scoring analysis against all 8 categories; gap identification', '2 weeks'],
        ['3. Custom Plan', 'Tailored infrastructure deployment plan with projected scoring', '1 week'],
        ['4. Agreement', 'Partnership signed; zero upfront cost', 'Week 4'],
        ['5. Deployment', 'Infrastructure deployed under society brand', 'Month 1-3'],
        ['6. Review', 'Annual performance review with plan refinement', 'Annual'],
    ])
doc2.add_page_break()

add_about(doc2)
doc2.add_page_break()

add_heading(doc2, 'Next Step', 1)
add_body(doc2, 'Your society has the expertise and the authority in your specialty. We have the infrastructure to transform that expertise into measurable, sustainable impact. Together, we build something that neither can build alone.')
add_body(doc2, '')
add_body(doc2, 'We invite your society\u2019s leadership to an initial consultation to explore how this partnership can accelerate your path to Category A.')
add_body(doc2, '')
add_contact(doc2, 'We build the infrastructure. You lead the specialty.')

path2 = r'C:\Users\د. ذيب\Desktop\Theeb-Ecosystem\05-SHARED\Brand-Assets\SCIENTIFIC_SOCIETY_PARTNERSHIP_PROPOSAL.docx'
doc2.save(path2)
print(f'Document 2 saved successfully!')
print(f'\nBoth proposals created at:\n{path1}\n{path2}')
